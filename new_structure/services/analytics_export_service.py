"""
Analytics Export Service for generating PDF, Word, and Excel reports.
Handles comprehensive analytics data export in multiple formats.
"""

import os
import io
import time
from datetime import datetime
from typing import Dict, Any, List, Optional
from flask import current_app
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import xlsxwriter
from ..services.academic_analytics_service import AcademicAnalyticsService


class AnalyticsExportService:
    """Service for exporting analytics data in various formats."""
    
    @classmethod
    def export_analytics_pdf(cls, analytics_data: Dict[str, Any], filters: Dict[str, Any] = None) -> bytes:
        """
        Export analytics data to PDF format.
        
        Args:
            analytics_data: Analytics data to export
            filters: Applied filters for context
            
        Returns:
            PDF file as bytes
        """
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1,  # Center alignment
                textColor=colors.HexColor('#2c5f5a')
            )
            story.append(Paragraph("School-Wide Academic Analytics Report", title_style))
            story.append(Spacer(1, 20))
            
            # Report metadata
            meta_style = ParagraphStyle(
                'MetaStyle',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#666666')
            )
            
            report_date = datetime.now().strftime("%B %d, %Y at %I:%M %p")
            story.append(Paragraph(f"Generated on: {report_date}", meta_style))
            
            if filters:
                filter_text = cls._format_filters_text(filters)
                story.append(Paragraph(f"Filters Applied: {filter_text}", meta_style))
            
            story.append(Spacer(1, 30))
            
            # Top Performers Section
            if analytics_data.get('top_performers'):
                story.append(Paragraph("Top Performing Students", styles['Heading2']))
                story.append(Spacer(1, 10))
                
                performers_table = cls._create_performers_table(analytics_data['top_performers'])
                story.append(performers_table)
                story.append(Spacer(1, 20))
            
            # Subject Performance Section
            if analytics_data.get('subject_analytics'):
                story.append(Paragraph("Subject Performance Analysis", styles['Heading2']))
                story.append(Spacer(1, 10))
                
                subjects_table = cls._create_subjects_table(analytics_data['subject_analytics'])
                story.append(subjects_table)
                story.append(Spacer(1, 20))
            
            # Summary Statistics
            if analytics_data.get('summary'):
                story.append(Paragraph("Summary Statistics", styles['Heading2']))
                story.append(Spacer(1, 10))
                
                summary_table = cls._create_summary_table(analytics_data['summary'])
                story.append(summary_table)
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            print(f"Error generating PDF: {e}")
            raise
    
    @classmethod
    def export_analytics_word(cls, analytics_data: Dict[str, Any], filters: Dict[str, Any] = None) -> bytes:
        """
        Export analytics data to Word format.
        
        Args:
            analytics_data: Analytics data to export
            filters: Applied filters for context
            
        Returns:
            Word document as bytes
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('School-Wide Academic Analytics Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Report metadata
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            
            if filters:
                filter_text = cls._format_filters_text(filters)
                doc.add_paragraph(f"Filters Applied: {filter_text}")
            
            doc.add_paragraph()  # Empty line
            
            # Top Performers Section
            if analytics_data.get('top_performers'):
                doc.add_heading('Top Performing Students', level=1)
                
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Rank'
                hdr_cells[1].text = 'Student Name'
                hdr_cells[2].text = 'Admission No.'
                hdr_cells[3].text = 'Average %'
                hdr_cells[4].text = 'Grade'
                
                # Data rows
                for performer in analytics_data['top_performers'][:10]:  # Top 10
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(performer.get('rank', ''))
                    row_cells[1].text = performer.get('name', '')
                    row_cells[2].text = performer.get('admission_number', '')
                    row_cells[3].text = f"{performer.get('average_percentage', 0):.1f}%"
                    row_cells[4].text = performer.get('grade_letter', '')
                
                doc.add_paragraph()  # Empty line
            
            # Subject Performance Section
            if analytics_data.get('subject_analytics'):
                doc.add_heading('Subject Performance Analysis', level=1)

                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Subject'
                hdr_cells[1].text = 'Teacher'
                hdr_cells[2].text = 'Average %'
                hdr_cells[3].text = 'Students'
                hdr_cells[4].text = 'Performance'

                # Data rows
                for subject in analytics_data['subject_analytics']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = subject.get('subject_name', '')
                    row_cells[1].text = subject.get('teacher_name', 'Not Assigned')
                    row_cells[2].text = f"{subject.get('average_percentage', 0):.1f}%"
                    row_cells[3].text = str(subject.get('student_count', 0))
                    row_cells[4].text = subject.get('performance_category', '')
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            print(f"Error generating Word document: {e}")
            raise
    
    @classmethod
    def export_analytics_excel(cls, analytics_data: Dict[str, Any], filters: Dict[str, Any] = None) -> bytes:
        """
        Export analytics data to Excel format.
        
        Args:
            analytics_data: Analytics data to export
            filters: Applied filters for context
            
        Returns:
            Excel file as bytes
        """
        try:
            buffer = io.BytesIO()
            workbook = xlsxwriter.Workbook(buffer, {'in_memory': True})
            
            # Define formats
            title_format = workbook.add_format({
                'bold': True,
                'font_size': 16,
                'align': 'center',
                'valign': 'vcenter',
                'bg_color': '#2c5f5a',
                'font_color': 'white'
            })
            
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#7dd3c0',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            data_format = workbook.add_format({
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            
            # Top Performers Sheet
            if analytics_data.get('top_performers'):
                worksheet1 = workbook.add_worksheet('Top Performers')
                
                # Title
                worksheet1.merge_range('A1:E1', 'Top Performing Students', title_format)
                worksheet1.set_row(0, 25)
                
                # Headers
                headers = ['Rank', 'Student Name', 'Admission No.', 'Average %', 'Grade']
                for col, header in enumerate(headers):
                    worksheet1.write(2, col, header, header_format)
                
                # Data
                for row, performer in enumerate(analytics_data['top_performers'], start=3):
                    worksheet1.write(row, 0, performer.get('rank', ''), data_format)
                    worksheet1.write(row, 1, performer.get('name', ''), data_format)
                    worksheet1.write(row, 2, performer.get('admission_number', ''), data_format)
                    worksheet1.write(row, 3, f"{performer.get('average_percentage', 0):.1f}%", data_format)
                    worksheet1.write(row, 4, performer.get('grade_letter', ''), data_format)
                
                # Auto-fit columns
                worksheet1.set_column('A:E', 15)
            
            # Subject Performance Sheet
            if analytics_data.get('subject_analytics'):
                worksheet2 = workbook.add_worksheet('Subject Performance')

                # Title
                worksheet2.merge_range('A1:E1', 'Subject Performance Analysis', title_format)
                worksheet2.set_row(0, 25)

                # Headers
                headers = ['Subject', 'Teacher', 'Average %', 'Students', 'Performance Category']
                for col, header in enumerate(headers):
                    worksheet2.write(2, col, header, header_format)

                # Data
                for row, subject in enumerate(analytics_data['subject_analytics'], start=3):
                    worksheet2.write(row, 0, subject.get('subject_name', ''), data_format)
                    worksheet2.write(row, 1, subject.get('teacher_name', 'Not Assigned'), data_format)
                    worksheet2.write(row, 2, f"{subject.get('average_percentage', 0):.1f}%", data_format)
                    worksheet2.write(row, 3, subject.get('student_count', 0), data_format)
                    worksheet2.write(row, 4, subject.get('performance_category', ''), data_format)

                # Auto-fit columns
                worksheet2.set_column('A:E', 20)
            
            # Summary Sheet
            if analytics_data.get('summary'):
                worksheet3 = workbook.add_worksheet('Summary')
                
                # Title
                worksheet3.merge_range('A1:B1', 'Analytics Summary', title_format)
                worksheet3.set_row(0, 25)
                
                summary = analytics_data['summary']
                row = 2
                
                summary_items = [
                    ('Total Students Analyzed', summary.get('total_students_analyzed', 0)),
                    ('Total Subjects Analyzed', summary.get('total_subjects_analyzed', 0)),
                    ('Report Generated', datetime.now().strftime('%B %d, %Y at %I:%M %p'))
                ]
                
                if filters:
                    summary_items.append(('Filters Applied', cls._format_filters_text(filters)))
                
                for item, value in summary_items:
                    worksheet3.write(row, 0, item, header_format)
                    worksheet3.write(row, 1, str(value), data_format)
                    row += 1
                
                worksheet3.set_column('A:B', 25)
            
            workbook.close()
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            print(f"Error generating Excel file: {e}")
            raise

    @staticmethod
    def _format_filters_text(filters: Dict[str, Any]) -> str:
        """Format filters dictionary into readable text."""
        filter_parts = []

        if filters.get('grade_id'):
            filter_parts.append(f"Grade: {filters.get('grade_name', filters['grade_id'])}")

        if filters.get('stream_id'):
            filter_parts.append(f"Stream: {filters.get('stream_name', filters['stream_id'])}")

        if filters.get('term_id'):
            filter_parts.append(f"Term: {filters.get('term_name', filters['term_id'])}")

        if filters.get('assessment_type_id'):
            filter_parts.append(f"Assessment: {filters.get('assessment_type_name', filters['assessment_type_id'])}")

        return ", ".join(filter_parts) if filter_parts else "None"

    @staticmethod
    def _create_performers_table(performers: List[Dict[str, Any]]) -> Table:
        """Create a table for top performers data."""
        data = [['Rank', 'Student Name', 'Admission No.', 'Average %', 'Grade']]

        for performer in performers[:10]:  # Top 10
            data.append([
                str(performer.get('rank', '')),
                performer.get('name', ''),
                performer.get('admission_number', ''),
                f"{performer.get('average_percentage', 0):.1f}%",
                performer.get('grade_letter', '')
            ])

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5f5a')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        return table

    @staticmethod
    def _create_subjects_table(subjects: List[Dict[str, Any]]) -> Table:
        """Create a table for subject performance data."""
        data = [['Subject', 'Teacher', 'Average %', 'Students', 'Performance Category']]

        for subject in subjects:
            data.append([
                subject.get('subject_name', ''),
                subject.get('teacher_name', 'Not Assigned'),
                f"{subject.get('average_percentage', 0):.1f}%",
                str(subject.get('student_count', 0)),
                subject.get('performance_category', '')
            ])

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7dd3c0')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        return table

    @staticmethod
    def _create_summary_table(summary: Dict[str, Any]) -> Table:
        """Create a table for summary statistics."""
        data = [
            ['Metric', 'Value'],
            ['Total Students Analyzed', str(summary.get('total_students_analyzed', 0))],
            ['Total Subjects Analyzed', str(summary.get('total_subjects_analyzed', 0))],
            ['Has Sufficient Data', 'Yes' if summary.get('has_sufficient_data', False) else 'No'],
            ['Report Generated', datetime.now().strftime('%B %d, %Y at %I:%M %p')]
        ]

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a9b8e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        return table
